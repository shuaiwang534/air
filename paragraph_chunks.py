from docx import Document
import re
import json
import os
import tempfile
from typing import Any, Dict, List, Optional, Tuple

from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

try:
    import pythoncom
    import win32com.client
except Exception:
    pythoncom = None
    win32com = None

INPUT_DOCX = "校验部分.doc"   # 这里现在也可以填 .doc
OUTPUT_JSONL = "section_chunks.jsonl"
OUTPUT_TABLES_JSONL = "output/tables_raw.jsonl"


def convert_doc_to_docx(src_path, dst_path):
    """
    把 .doc 转成 .docx
    需要：Windows + 已安装 Microsoft Word + pywin32
    """
    if pythoncom is None or win32com is None:
        raise RuntimeError("处理 .doc 需要安装 pywin32，并确保本机安装了 Microsoft Word。")

    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word_doc = None

    try:
        word_doc = word.Documents.Open(os.path.abspath(src_path), ReadOnly=True)
        # wdFormatXMLDocument = 16  -> .docx
        word_doc.SaveAs(os.path.abspath(dst_path), FileFormat=16)
    finally:
        if word_doc is not None:
            word_doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()


def load_document_any(input_path):
    """
    同时支持 .doc / .docx
    返回: (doc对象, 临时docx路径或None)
    """
    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".docx":
        return Document(input_path), None

    if ext == ".doc":
        fd, tmp_docx = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        if os.path.exists(tmp_docx):
            try:
                os.remove(tmp_docx)
            except Exception:
                pass

        convert_doc_to_docx(input_path, tmp_docx)
        return Document(tmp_docx), tmp_docx

    raise ValueError("只支持 .doc 和 .docx 文件：{0}".format(input_path))


def _parse_outline_val(vals):
    if not vals:
        return None
    try:
        level = int(vals[0]) + 1   # Word里 0->1级目录, 1->2级目录
        if 1 <= level <= 9:
            return level
    except Exception:
        pass
    return None


def _get_style_heading_level(style, visited=None):
    """
    从样式及其继承链中读取标题级别
    兼容：
    1. Heading 1 / Heading 2 ...
    2. 标题 1 / 标题 2 ...
    3. 样式本身带 outlineLvl
    """
    if style is None:
        return None

    if visited is None:
        visited = set()

    sid = id(style)
    if sid in visited:
        return None
    visited.add(sid)

    # 1) 样式XML中的目录级别
    try:
        vals = style.element.xpath('./w:pPr/w:outlineLvl/@w:val')
        lvl = _parse_outline_val(vals)
        if lvl is not None:
            return lvl
    except Exception:
        pass

    # 2) 样式名匹配
    try:
        name = style.name.strip()
    except Exception:
        name = ""

    m = re.match(r'^(?:Heading|标题)\s*([1-9])$', name, re.IGNORECASE)
    if m:
        return int(m.group(1))

    # 3) 递归查父样式
    try:
        base_style = style.base_style
    except Exception:
        base_style = None

    if base_style is not None:
        return _get_style_heading_level(base_style, visited)

    return None


def get_heading_level(p):
    """
    优先读取段落本身的“目录级别”
    如果没有，再回退到样式/父样式
    """
    # 1) 段落本身的目录级别
    try:
        vals = p._p.xpath('./w:pPr/w:outlineLvl/@w:val')
        lvl = _parse_outline_val(vals)
        if lvl is not None:
            return lvl
    except Exception:
        pass

    # 2) 样式中的目录级别/Heading样式
    try:
        return _get_style_heading_level(p.style)
    except Exception:
        return None


def parse_heading(text):
    """
    解析以下形式：
    - 8.7.1液压余度配置
    - 8.7.1 液压余度配置
    - 8.7.1：液压余度配置
    - 8.7.1-液压余度配置
    - 8.7.1 液压余度配置（主飞控）
    """
    text = text.strip()

    m = re.match(
        r"""
        ^(\d+(?:\.\d+)*)        # 章节编号
        [\s:：\-—]*             # 可选分隔符（空格/中文冒号/连字符）
        (.+?)                   # 标题正文
        $
        """,
        text,
        re.VERBOSE
    )

    if m:
        return m.group(1), m.group(2).strip()

    return None, text


def _clean_cell_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n+", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_table_rows(table: Table) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        rows.append([_clean_cell_text(cell.text) for cell in row.cells])
    return rows


def extract_table_text(raw_rows: List[List[str]]) -> str:
    """提取表格内容, 转为文本表示"""
    lines = [" | ".join(row).strip() for row in raw_rows]
    return "\n".join([line for line in lines if line])


def build_table_record(
    table: Table,
    table_id: str,
    current_section: Optional[Dict[str, Any]],
) -> Dict[str, Any]:
    raw_rows = extract_table_rows(table)
    header = raw_rows[0] if raw_rows else []
    rows = raw_rows[1:] if len(raw_rows) > 1 else []
    raw_text = extract_table_text(raw_rows)

    section_id = ""
    title = ""
    path: List[str] = []
    if current_section:
        section_id = current_section.get("section_id") or ""
        title = current_section.get("title") or ""
        path = current_section.get("path") or []

    return {
        "table_id": table_id,
        "section_id": section_id,
        "title": title,
        "path": path,
        "header": header,
        "rows": rows,
        "raw_rows": raw_rows,
        "raw_text": raw_text,
    }


def iter_block_items(doc):
    """
    按文档真实顺序遍历 段落 和 表格
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, doc)


# =========================
# 前处理：页脚 / 目录 / 图片 / 图题
# =========================
def clear_all_footers(doc):
    """
    保险处理：清空所有 section 的页脚文本。
    注意：你当前的正文遍历本来就不会读到页脚，这里只是额外清理。
    """
    for section in doc.sections:
        footer_names = [
            "footer",
            "first_page_footer",
            "even_page_footer",
        ]
        for name in footer_names:
            try:
                footer = getattr(section, name)
            except Exception:
                footer = None

            if footer is None:
                continue

            try:
                for p in footer.paragraphs:
                    p.text = ""
            except Exception:
                pass


def paragraph_has_image(p):
    """
    判断段落里是否含图片
    支持常见的 w:drawing / w:pict
    """
    try:
        if p._p.xpath('.//w:drawing'):
            return True
        if p._p.xpath('.//w:pict'):
            return True
    except Exception:
        pass
    return False


def is_figure_caption(p):
    """
    判断段落是否像“图片小标题 / 图题”
    只删图片对应的小标题，不删表题
    """
    text = p.text.strip()
    if not text:
        return False

    try:
        style_name = p.style.name.strip()
    except Exception:
        style_name = ""

    if re.match(r'^(?:Caption|题注)$', style_name, re.IGNORECASE):
        if re.match(
            r'^\s*(?:图|Figure|FIGURE|Fig\.?|FIG\.?)\s*'
            r'[A-Za-z0-9一二三四五六七八九十百千IVXivx\-\.\(\)（）]+'
            r'\s*[:：\-—]?\s*\S+',
            text,
            re.IGNORECASE
        ):
            return True

    if re.match(
        r'^\s*(?:图|Figure|FIGURE|Fig\.?|FIG\.?)\s*'
        r'[A-Za-z0-9一二三四五六七八九十百千IVXivx\-\.\(\)（）]+'
        r'\s*[:：\-—]?\s*\S+',
        text,
        re.IGNORECASE
    ):
        return True

    return False


def _remove_paragraph(p):
    """
    从 XML 中删除段落
    """
    try:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    except Exception:
        pass


def is_toc_paragraph(p):
    """
    判断是否为目录段落
    """
    text = p.text.strip()

    try:
        style_name = p.style.name.strip()
    except Exception:
        style_name = ""

    # 常见目录样式：TOC 1 / TOC 2 ...
    if re.match(r'^(?:TOC|目录)\s*[1-9]$', style_name, re.IGNORECASE):
        return True

    # 目录标题
    if re.match(r'^(?:目录|Contents?)$', text, re.IGNORECASE):
        return True

    # Word 自动目录域
    try:
        xml = p._p.xml
        if ' TOC ' in xml or 'TOC \\\\o' in xml or 'TOC \\\\h' in xml:
            return True
    except Exception:
        pass

    # 常见目录行：xxxx........12
    if re.match(r'^.+\.{2,}\s*\d+\s*$', text):
        return True

    return False


def remove_toc(doc):
    """
    删除文档开头的目录部分
    逻辑：
    - 从前往后找，遇到“目录/TOC样式/TOC域”开始进入目录区
    - 连续删除目录区段落
    - 遇到第一个真正正文标题时停止
    """
    body_paragraphs = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            body_paragraphs.append(Paragraph(child, doc))

    to_remove = []
    in_toc = False

    for p in body_paragraphs:
        text = p.text.strip()

        if not in_toc:
            if is_toc_paragraph(p):
                in_toc = True
                to_remove.append(p)
            continue

        # 已进入目录区
        if not text:
            to_remove.append(p)
            continue

        if is_toc_paragraph(p):
            to_remove.append(p)
            continue

        # 遇到真正标题 -> 目录结束
        lvl = get_heading_level(p)
        if lvl is not None:
            break

        # 有些目录项不是 TOC 样式，但表现为“文字 + 页码”
        if re.match(r'^.+\s+\d+\s*$', text):
            to_remove.append(p)
            continue

        break

    for p in to_remove:
        _remove_paragraph(p)


def remove_images_and_captions(doc):
    """
    删除正文中的图片段落，以及图片前后紧邻的图题段落
    不删除表格
    """
    body_paragraphs = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            body_paragraphs.append(Paragraph(child, doc))

    to_remove = []

    n = len(body_paragraphs)
    for i, p in enumerate(body_paragraphs):
        if not paragraph_has_image(p):
            continue

        # 1. 删除图片所在段落
        to_remove.append(p)

        # 2. 删除后一个非空段落（如果像图题）
        j = i + 1
        while j < n and not body_paragraphs[j].text.strip():
            j += 1
        if j < n and is_figure_caption(body_paragraphs[j]):
            to_remove.append(body_paragraphs[j])

        # 3. 删除前一个非空段落（如果像图题）
        j = i - 1
        while j >= 0 and not body_paragraphs[j].text.strip():
            j -= 1
        if j >= 0 and is_figure_caption(body_paragraphs[j]):
            to_remove.append(body_paragraphs[j])

    seen = set()
    unique_remove = []
    for p in to_remove:
        pid = id(p._element)
        if pid not in seen:
            seen.add(pid)
            unique_remove.append(p)

    for p in unique_remove:
        _remove_paragraph(p)


def preprocess_document(doc):
    """
    文档前处理：
    1. 清空页脚
    2. 删除开头目录
    3. 删除正文中的图片及其对应小标题
    """
    clear_all_footers(doc)
    remove_toc(doc)
    remove_images_and_captions(doc)


# =========================
# 正文切分
# =========================
def build_section_chunks(doc) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    chunks: List[Dict[str, Any]] = []
    table_records: List[Dict[str, Any]] = []
    heading_stack = []   # [{"level": 1, "title": "..."}]
    current = None
    table_index = 0

    for elem_type, elem in iter_block_items(doc):
        if elem_type == "paragraph":
            text = elem.text.strip()
            if not text:
                continue

            lvl = get_heading_level(elem)

            if lvl is not None:
                section_id, title = parse_heading(text)

                while heading_stack and heading_stack[-1]["level"] >= lvl:
                    heading_stack.pop()

                heading_stack.append({
                    "level": lvl,
                    "title": title
                })

                current = {
                    "section_id": section_id,
                    "title": title,
                    "path": [h["title"] for h in heading_stack],
                    "content_lines": []
                }
                chunks.append(current)
                continue

            if current:
                current["content_lines"].append(text)

        elif elem_type == "table":
            if not current:
                continue

            table_index += 1
            section_id = current.get("section_id") or "NO_SECTION"
            table_id = "{0}-T{1:03d}".format(section_id, table_index)

            table_record = build_table_record(
                table=elem,
                table_id=table_id,
                current_section=current,
            )
            table_records.append(table_record)
            current["content_lines"].append("[表格引用:{0}]".format(table_id))

    for c in chunks:
        c["content"] = "\n".join(c["content_lines"]).strip()
        del c["content_lines"]

    return chunks, table_records


def main():
    doc = None
    temp_docx = None

    try:
        doc, temp_docx = load_document_any(INPUT_DOCX)

        preprocess_document(doc)
        chunks, table_records = build_section_chunks(doc)

        out_dir = os.path.dirname(OUTPUT_JSONL) or "."
        os.makedirs(out_dir, exist_ok=True)
        with open(OUTPUT_JSONL, "w", encoding="utf-8") as f:
            for c in chunks:
                f.write(json.dumps(c, ensure_ascii=False) + "\n")

        table_dir = os.path.dirname(OUTPUT_TABLES_JSONL) or "."
        os.makedirs(table_dir, exist_ok=True)
        with open(OUTPUT_TABLES_JSONL, "w", encoding="utf-8") as f:
            for t in table_records:
                f.write(json.dumps(t, ensure_ascii=False) + "\n")

        print("生成章节数：{0}".format(len(chunks)))
        print("生成表格数：{0}".format(len(table_records)))

    finally:
        if temp_docx and os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except Exception:
                pass


if __name__ == "__main__":
    main()
