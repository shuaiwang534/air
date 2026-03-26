# -*- coding: utf-8 -*-

from docx import Document
import re
import json
import os
import tempfile
from collections import OrderedDict
from typing import Any, Dict, List, Optional, Tuple

from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

try:
    import pythoncom
    import win32com.client
except Exception:
    pythoncom = None
    win32com = None

INPUT_DOCX = "校验部分.doc"   # 支持 .doc 或 .docx
OUTPUT_JSONL = "section_chunks.jsonl"
OUTPUT_TABLES_JSONL = "output/table_rows.jsonl"


def resolve_input_doc_path(input_path):
    """
    Resolve input path with .doc/.docx compatibility:
    1) direct path
    2) same-stem sibling with swapped extension
    3) first .docx/.doc in current directory
    """
    if os.path.exists(input_path):
        return input_path

    abs_input = os.path.abspath(input_path)
    root, ext = os.path.splitext(abs_input)
    ext = ext.lower()

    if ext in (".doc", ".docx"):
        alt = root + (".doc" if ext == ".docx" else ".docx")
        if os.path.exists(alt):
            return alt

    cwd = os.getcwd()
    for name in os.listdir(cwd):
        low = name.lower()
        if low.endswith(".docx"):
            return os.path.join(cwd, name)
    for name in os.listdir(cwd):
        low = name.lower()
        if low.endswith(".doc"):
            return os.path.join(cwd, name)

    raise FileNotFoundError("Input document not found: {0}".format(abs_input))


def convert_doc_to_docx(src_path, dst_path):
    """
    灏?.doc 杞垚 .docx
    瑕佹眰锛歐indows + 宸插畨瑁?Microsoft Word + pywin32
    """
    if pythoncom is None or win32com is None:
        raise RuntimeError("Processing .doc requires pywin32 and Microsoft Word installed.")

    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    word_doc = None

    try:
        word_doc = word.Documents.Open(os.path.abspath(src_path), ReadOnly=True)
        # wdFormatXMLDocument = 16  -> .docx
        word_doc.SaveAs(os.path.abspath(dst_path), FileFormat=16)
    finally:
        if word_doc is not None:
            word_doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()


def load_document_any(input_path):
    """
    鍚屾椂鏀寔 .doc / .docx
    杩斿洖: (doc瀵硅薄, 涓存椂docx璺緞鎴朜one)
    """
    resolved_input = resolve_input_doc_path(input_path)

    ext = os.path.splitext(resolved_input)[1].lower()

    if ext == ".docx":
        return Document(resolved_input), None

    if ext == ".doc":
        fd, tmp_docx = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        if os.path.exists(tmp_docx):
            try:
                os.remove(tmp_docx)
            except Exception:
                pass

        try:
            convert_doc_to_docx(resolved_input, tmp_docx)
        except Exception as e:
            sibling_docx = os.path.splitext(resolved_input)[0] + ".docx"
            if os.path.exists(sibling_docx):
                return Document(sibling_docx), None
            raise RuntimeError(
                "Failed to convert .doc to .docx: {0}\n"
                "If you are using WPS, please open the file and save it as .docx, then rerun with --input-doc."
                .format(e)
            )
        return Document(tmp_docx), tmp_docx

    raise ValueError("浠呮敮鎸?.doc 鎴?.docx 鏂囦欢: {0}".format(resolved_input))


def _parse_outline_val(vals):
    if not vals:
        return None
    try:
        level = int(vals[0]) + 1
        if 1 <= level <= 9:
            return level
    except Exception:
        pass
    return None


def _get_style_heading_level(style, visited=None):
    """
    浠庢牱寮忎互鍙婂叾缁ф壙閾捐鍙栨爣棰樼骇鍒€?
    鍏煎:
    1) Heading 1 / Heading 2 ...
    2) 鏍囬 1 / 鏍囬 2 ...
    3) 鏍峰紡鑷韩甯?outlineLvl
    """
    if style is None:
        return None

    if visited is None:
        visited = set()

    sid = id(style)
    if sid in visited:
        return None
    visited.add(sid)

    # 1) 鏍峰紡 XML 涓殑 outline level
    try:
        vals = style.element.xpath('./w:pPr/w:outlineLvl/@w:val')
        lvl = _parse_outline_val(vals)
        if lvl is not None:
            return lvl
    except Exception:
        pass

    # 2) 鏍峰紡鍚嶅尮閰?
    try:
        name = style.name.strip()
    except Exception:
        name = ""

    m = re.match(r'^(?:Heading|鏍囬)\s*([1-9])$', name, re.IGNORECASE)
    if m:
        return int(m.group(1))

    # 3) 閫掑綊妫€鏌ョ埗鏍峰紡
    try:
        base_style = style.base_style
    except Exception:
        base_style = None

    if base_style is not None:
        return _get_style_heading_level(base_style, visited)

    return None


def get_heading_level(p):
    """
    浼樺厛璇诲彇娈佃惤鑷韩 outline level锛涘鏋滄病鏈夊垯鍥為€€鍒版牱寮忛摼銆?
    """
    try:
        vals = p._p.xpath('./w:pPr/w:outlineLvl/@w:val')
        lvl = _parse_outline_val(vals)
        if lvl is not None:
            return lvl
    except Exception:
        pass

    try:
        return _get_style_heading_level(p.style)
    except Exception:
        return None


def parse_heading(text):
    """
    Parse heading lines like:
    - 8.7.1液压余度配置
    - 8.7.1 液压余度配置
    - 8.7.1：液压余度配置
    - 8.7.1-液压余度配置
    """
    text = text.strip()

    m = re.match(r"^(\d+(?:\.\d+)*)(?:\s*[:：\-—]\s*|\s+)?(.+)$", text)

    if m:
        return m.group(1), m.group(2).strip()

    return None, text


def _parse_numeric_section_id(section_id):
    text = str(section_id or "").strip()
    if not text:
        return []

    parts = text.split(".")
    nums = []
    for p in parts:
        if not p.isdigit():
            return []
        nums.append(int(p))
    return nums


def _sync_section_counters(counters, section_id):
    nums = _parse_numeric_section_id(section_id)
    if not nums:
        return

    counters[:] = nums


def _gen_section_id_from_level(level, counters):
    lvl = int(level or 0)
    if lvl <= 0:
        return ""

    while len(counters) < lvl:
        counters.append(0)

    del counters[lvl:]

    # If parent levels are missing, initialize them to 1.
    for i in range(lvl - 1):
        if counters[i] <= 0:
            counters[i] = 1

    counters[lvl - 1] += 1
    return ".".join([str(x) for x in counters])


def iter_block_items(doc):
    """
    鎸夋枃妗ｇ湡瀹為『搴忛亶鍘嗘钀藉拰琛ㄦ牸銆?
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, doc)


def _clean_table_text(value):
    if value is None:
        return ""
    text = str(value).replace("\r\n", "\n").replace("\r", "\n")
    parts = [x.strip() for x in text.split("\n") if x.strip()]
    return " ".join(parts).strip()


def _pad_rows(rows):
    width = 0
    for row in rows:
        if len(row) > width:
            width = len(row)
    out = []
    for row in rows:
        out.append(list(row) + [""] * (width - len(row)))
    return out


def _collapse_repeated_long_cells(cells, min_len=16):
    """
    Handle merged-like horizontal repeats from some Word tables:
    if adjacent columns contain exactly the same long text, keep the first one
    and clear the trailing duplicates.
    """
    out = list(cells)
    n = len(out)
    i = 0
    while i < n:
        cur = _clean_table_text(out[i])
        if not cur:
            i += 1
            continue

        j = i + 1
        while j < n and _clean_table_text(out[j]) == cur:
            j += 1

        if (j - i) >= 2 and len(cur) >= int(min_len):
            for k in range(i + 1, j):
                out[k] = ""

        i = j
    return out


def _normalize_table_rows(table):
    rows = []
    for row in table.rows:
        cells = [_clean_table_text(cell.text) for cell in row.cells]
        rows.append(cells)
    return _pad_rows(rows)


def _fill_right(rows, upto_row_count):
    out = [list(row) for row in rows]
    upto = min(upto_row_count, len(out))
    for i in range(upto):
        last = ""
        for j in range(len(out[i])):
            val = _clean_table_text(out[i][j])
            if val:
                last = val
            elif last:
                out[i][j] = last
    return out


def _fill_down(rows, start_row):
    out = [list(row) for row in rows]
    if not out:
        return out
    width = len(out[0])
    for j in range(width):
        last = ""
        for i in range(start_row, len(out)):
            val = _clean_table_text(out[i][j])
            if val:
                last = val
            elif last:
                out[i][j] = last
    return out


def _build_header_from_first_row(first_row):
    headers = []
    seen = {}
    for i, cell in enumerate(first_row):
        name = _clean_table_text(cell) or "col_{0}".format(i + 1)
        count = seen.get(name, 0) + 1
        seen[name] = count
        if count > 1:
            name = "{0}__dup{1}".format(name, count)
        headers.append(name)
    return headers


def build_table_row_records(table, table_id, section_context):
    """
    Build row-wise JSON records:
    - first row as header
    - second row onwards as data rows
    - fill-right on header row + fill-down on data rows
    """
    raw_rows = _normalize_table_rows(table)
    if not raw_rows:
        return []

    header_row_count = 1
    rows_after_right = _fill_right(raw_rows, header_row_count)
    rows_after_down = _fill_down(rows_after_right, header_row_count)

    header = _build_header_from_first_row(rows_after_down[0])
    data_rows = rows_after_down[1:]

    section_id = section_context.get("section_id") or ""
    title = section_context.get("title") or ""
    path = section_context.get("path") or []

    records = []
    row_no = 0
    for row in data_rows:
        if not any(_clean_table_text(x) for x in row):
            continue

        cells = _collapse_repeated_long_cells(list(row))
        local_header = list(header)

        if len(cells) < len(local_header):
            cells = cells + [""] * (len(local_header) - len(cells))
        elif len(cells) > len(local_header):
            for extra_idx in range(len(local_header), len(cells)):
                local_header.append("col_{0}".format(extra_idx + 1))

        row_no += 1
        row_map = OrderedDict()
        for k, v in zip(local_header, cells):
            row_map[k] = _clean_table_text(v)

        records.append(
            {
                "table_id": table_id,
                "section_id": section_id,
                "title": title,
                "path": path,
                "row_index": row_no,
                "header": local_header,
                "cells": cells,
                "row_map": row_map,
            }
        )

    return records


# =========================
# 鍓嶅鐞嗭細椤佃剼 / 鐩綍 / 鍥剧墖 / 鍥鹃
# =========================
def clear_all_footers(doc):
    """
    淇濋櫓澶勭悊锛氭竻绌烘墍鏈?section 鐨勯〉鑴氭枃鏈€?
    娉細褰撳墠姝ｆ枃閬嶅巻鏈潵灏变笉浼氳鍒伴〉鑴氾紝杩欓噷鍙槸棰濆娓呯悊銆?
    """
    for section in doc.sections:
        footer_names = [
            "footer",
            "first_page_footer",
            "even_page_footer",
        ]
        for name in footer_names:
            try:
                footer = getattr(section, name)
            except Exception:
                footer = None

            if footer is None:
                continue

            try:
                for p in footer.paragraphs:
                    p.text = ""
            except Exception:
                pass


def paragraph_has_image(p):
    """
    鍒ゆ柇娈佃惤涓槸鍚﹀寘鍚浘鐗囥€?
    鏀寔甯歌 w:drawing / w:pict銆?
    """
    try:
        if p._p.xpath('.//w:drawing'):
            return True
        if p._p.xpath('.//w:pict'):
            return True
    except Exception:
        pass
    return False


def is_figure_caption(p):
    """
    判断段落是否像“图片小标题/图题”。
    只删除图片对应的图题，不删除表题。
    """
    text = p.text.strip()
    if not text:
        return False

    try:
        style_name = p.style.name.strip()
    except Exception:
        style_name = ""

    fig_caption_pattern = re.compile(
        r'^\s*(?:图|Figure|FIGURE|Fig\.?|FIG\.?)\s*'
        r'[A-Za-z0-9一二三四五六七八九十百千万IVXivx\-\.\(\)（）]+'
        r'\s*(?:[:：\-—]\s*)?\S+',
        re.IGNORECASE,
    )

    # style 命中 caption 时更宽松，但仍要求是图题形态
    if re.match(r'^(?:Caption|题注)$', style_name, re.IGNORECASE):
        if fig_caption_pattern.match(text):
            return True

    if fig_caption_pattern.match(text):
        return True

    return False


def _remove_paragraph(p):
    """
    浠?XML 涓垹闄ゆ钀姐€?
    """
    try:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)
    except Exception:
        pass


def is_toc_paragraph(p):
    """
    鍒ゆ柇娈佃惤鏄惁鏄洰褰曢」銆?
    """
    text = p.text.strip()

    try:
        style_name = p.style.name.strip()
    except Exception:
        style_name = ""

    if re.match(r'^(?:TOC|鐩綍)\s*[1-9]$', style_name, re.IGNORECASE):
        return True
    if re.match(r'^(?:鐩綍|Contents?)$', text, re.IGNORECASE):
        return True

    try:
        xml = p._p.xml
        if " TOC " in xml or "TOC \\\\o" in xml or "TOC \\\\h" in xml:
            return True
    except Exception:
        pass

    if re.match(r'^.+\.{2,}\s*\d+\s*$', text):
        return True

    return False


def remove_toc(doc):
    """
    鍒犻櫎鏂囨。寮€澶寸洰褰曞尯鍩熴€?
    瑙勫垯锛?
    1. 鍙戠幇鐩綍娈佃惤鍚庤繘鍏ョ洰褰曞尯銆?
    2. 杩炵画鍒犻櫎鐩綍鍖烘钀斤紙鍚┖琛?鐩綍鏍峰紡/鐩綍琛岋級銆?
    3. 閬囧埌绗竴涓湡瀹炴鏂囨爣棰樺悗鍋滄銆?
    """
    body_paragraphs = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            body_paragraphs.append(Paragraph(child, doc))

    to_remove = []
    in_toc = False

    for p in body_paragraphs:
        text = p.text.strip()

        if not in_toc:
            if is_toc_paragraph(p):
                in_toc = True
                to_remove.append(p)
            continue

        # 宸茶繘鍏ョ洰褰曞尯
        if not text:
            to_remove.append(p)
            continue

        if is_toc_paragraph(p):
            to_remove.append(p)
            continue

        # 閬囧埌鐪熷疄鏍囬锛岀洰褰曠粨鏉?
        lvl = get_heading_level(p)
        if lvl is not None:
            break

        # 鏈変簺鐩綍琛屼笉鏄?TOC 鏍峰紡锛屼絾褰㈡€佹槸鈥滄爣棰?+ 椤电爜鈥?
        if re.match(r'^.+\s+\d+\s*$', text):
            to_remove.append(p)
            continue

        break

    for p in to_remove:
        _remove_paragraph(p)


def remove_images_and_captions(doc):
    """
    鍒犻櫎姝ｆ枃涓殑鍥剧墖娈佃惤锛屼互鍙婂浘鐗囧墠鍚庣揣閭荤殑鍥鹃娈佃惤銆?
    涓嶅垹闄よ〃鏍笺€?
    """
    body_paragraphs = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            body_paragraphs.append(Paragraph(child, doc))

    to_remove = []

    n = len(body_paragraphs)
    for i, p in enumerate(body_paragraphs):
        if not paragraph_has_image(p):
            continue

        # 1) 鍒犻櫎鍥剧墖鎵€鍦ㄦ钀?
        to_remove.append(p)

        # 2) 鍒犻櫎鍚庝竴涓潪绌烘钀斤紙鑻ュ儚鍥鹃锛?
        j = i + 1
        while j < n and not body_paragraphs[j].text.strip():
            j += 1
        if j < n and is_figure_caption(body_paragraphs[j]):
            to_remove.append(body_paragraphs[j])

        # 3) 鍒犻櫎鍓嶄竴涓潪绌烘钀斤紙鑻ュ儚鍥鹃锛?
        j = i - 1
        while j >= 0 and not body_paragraphs[j].text.strip():
            j -= 1
        if j >= 0 and is_figure_caption(body_paragraphs[j]):
            to_remove.append(body_paragraphs[j])

    seen = set()
    unique_remove = []
    for p in to_remove:
        pid = id(p._element)
        if pid not in seen:
            seen.add(pid)
            unique_remove.append(p)

    for p in unique_remove:
        _remove_paragraph(p)


def preprocess_document(doc):
    """
    鏂囨。棰勫鐞嗭細
    1) 娓呯┖椤佃剼
    2) 鍒犻櫎寮€澶寸洰褰?
    3) 鍒犻櫎姝ｆ枃鍥剧墖鍙婂浘棰?
    """
    clear_all_footers(doc)
    remove_toc(doc)
    remove_images_and_captions(doc)


# =========================
# 姝ｆ枃鍒囧垎
# =========================
def build_section_chunks(doc) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    chunks: List[Dict[str, Any]] = []
    table_row_records: List[Dict[str, Any]] = []
    heading_stack = []   # [{"level": 1, "title": "..."}]
    section_counters: List[int] = []
    current = None
    table_index = 0

    for elem_type, elem in iter_block_items(doc):
        if elem_type == "paragraph":
            text = elem.text.strip()
            if not text:
                continue

            lvl = get_heading_level(elem)

            if lvl is not None:
                section_id, title = parse_heading(text)
                if section_id:
                    _sync_section_counters(section_counters, section_id)
                else:
                    section_id = _gen_section_id_from_level(lvl, section_counters)

                while heading_stack and heading_stack[-1]["level"] >= lvl:
                    heading_stack.pop()

                heading_stack.append({
                    "level": lvl,
                    "title": title
                })

                current = {
                    "section_id": section_id or "",
                    "title": title,
                    "path": [h["title"] for h in heading_stack],
                    "content_lines": []
                }
                chunks.append(current)
                continue

            if current:
                current["content_lines"].append(text)

        elif elem_type == "table":
            table_index += 1

            table_context = current
            if table_context is None:
                # 表格可能出现在首个标题之前，保留结构而不直接丢弃
                table_context = {
                    "section_id": "",
                    "title": heading_stack[-1]["title"] if heading_stack else "",
                    "path": [h["title"] for h in heading_stack] if heading_stack else [],
                }

            section_id = table_context.get("section_id") or "NO_SECTION"
            table_id = "{0}-T{1:03d}".format(section_id, table_index)

            row_records = build_table_row_records(
                table=elem,
                table_id=table_id,
                section_context=table_context,
            )
            table_row_records.extend(row_records)

            if current:
                current["content_lines"].append("[表格引用:{0}]".format(table_id))

    for c in chunks:
        c["content"] = "\n".join(c["content_lines"]).strip()
        del c["content_lines"]

    return chunks, table_row_records


def main():
    doc = None
    temp_docx = None

    try:
        doc, temp_docx = load_document_any(INPUT_DOCX)

        preprocess_document(doc)
        chunks, table_row_records = build_section_chunks(doc)

        out_dir = os.path.dirname(OUTPUT_JSONL) or "."
        os.makedirs(out_dir, exist_ok=True)
        with open(OUTPUT_JSONL, "w", encoding="utf-8") as f:
            for c in chunks:
                f.write(json.dumps(c, ensure_ascii=False) + "\n")

        table_dir = os.path.dirname(OUTPUT_TABLES_JSONL) or "."
        os.makedirs(table_dir, exist_ok=True)
        with open(OUTPUT_TABLES_JSONL, "w", encoding="utf-8") as f:
            for t in table_row_records:
                f.write(json.dumps(t, ensure_ascii=False) + "\n")

        print("生成章节数：{0}".format(len(chunks)))
        unique_table_ids = set([x.get("table_id") for x in table_row_records if x.get("table_id")])
        print("生成表格数：{0}".format(len(unique_table_ids)))
        print("生成表格行数：{0}".format(len(table_row_records)))
        if not table_row_records:
            print("[WARN] 未提取到任何表格数据行，请检查输入文档是否包含正文表格。")

    finally:
        if temp_docx and os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except Exception:
                pass


if __name__ == "__main__":
    main()



